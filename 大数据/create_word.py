from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# 设置字体
def set_run_font(run, font_name='微软雅黑', size=11, bold=False, color=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

# 添加标题
title = doc.add_heading('苹果蓝牙耳机售后AI聊天机器人项目分工表', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 课程信息
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('课程信息')
set_run_font(run, size=14, bold=True)

info_items = [
    '课程名称：大数据机器学习III',
    '项目周期：10周',
    '团队人数：3人'
]
for item in info_items:
    p = doc.add_paragraph(item, style='List Bullet')

# 团队成员
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('团队成员角色分配')
set_run_font(run, size=14, bold=True)

members = [
    '成员A（jiuyong）：项目经理 / 后端负责人',
    '成员B（a）：前端负责人 / 测试工程师',
    '成员C（b）：文档负责人 / 数据库 / 运维'
]
for member in members:
    p = doc.add_paragraph(member, style='List Bullet')

# 阶段数据
stages = [
    {
        'title': '第一阶段：市场调研与项目定位（第1-2周）',
        'tasks': [
            ['市场调研统筹', 'jiuyong', '牵头组织市场调研，分析苹果官方售后流程和痛点，确定项目差异化定位', '调研方向文档'],
            ['竞品分析', 'a', '调研现有AI客服机器人的技术方案和用户体验，收集竞品信息分析优劣势', '竞品分析报告'],
            ['用户画像定义', 'a', '协助定义目标用户画像（电商消费者、线下门店顾客等）', '用户画像文档'],
            ['报告撰写', 'b', '整理调研数据，撰写"项目背景与市场分析"章节，明确核心功能清单（产品咨询、订单查询、退换货政策解答、故障排查引导、客户情绪安抚）', '市场分析摘要'],
            ['检查点', '全员', '提交市场分析摘要，团队评审', '市场分析摘要定稿']
        ]
    },
    {
        'title': '第二阶段：技术选型与系统设计（第2-3周）',
        'tasks': [
            ['技术选型', 'jiuyong', '主导技术栈选型（前端：Vue/React，后端：Node.js/Python，数据库：MySQL/MongoDB，AI：DeepSeek/通义千问）', '技术选型文档'],
            ['系统架构设计', 'jiuyong', '设计系统架构，绘制技术架构图（前端、后端、数据库、AI模型服务），设计API接口规范', '系统架构图、API文档'],
            ['UI原型设计', 'a', '设计用户聊天界面原型，设计管理员后台界面', 'UI原型图'],
            ['数据库设计', 'b', '设计数据库ER图，规划数据表结构', 'ER图'],
            ['报告撰写', 'b', '撰写"系统架构设计"章节，编写系统设计文档初稿', '系统设计文档'],
            ['检查点', '全员', '提交系统设计文档初稿，团队评审', '系统设计文档定稿']
        ]
    },
    {
        'title': '第三阶段：AI功能核心实现（第3-5周）',
        'tasks': [
            ['大模型对接', 'jiuyong', '精准对接大模型API（DeepSeek/通义千问等），实现对话接口，处理用户输入和AI响应', 'AI调用接口代码'],
            ['提示词工程', 'jiuyong', '设计并优化系统提示词，确保回答专业精准，具备边界意识（无法回答时的引导策略）', '提示词文档'],
            ['知识库构建', 'a', '创建苹果蓝牙耳机销售/售后知识库（Q&A对、产品手册摘要、退换货政策、常见故障排查）', '知识库数据'],
            ['AI测试', 'a', '测试AI回答的准确性和专业性，反馈优化建议', 'AI测试报告'],
            ['文档撰写', 'b', '记录提示词设计思路和调优过程，撰写"AI模型集成与优化"章节，整理知识库文档', 'AI集成文档'],
            ['检查点', '全员', '完成基础对话接口，能处理用户输入并返回AI生成的回答', '可对话的AI原型']
        ]
    },
    {
        'title': '第四阶段：全栈开发与实现（第5-7周）',
        'tasks': [
            ['后端API开发', 'jiuyong', '实现用户会话管理、知识库CRUD接口、AI模型调用代理、对话历史存储等API', '后端API代码'],
            ['数据库实现', 'jiuyong', '根据ER图创建数据库，实现数据持久化', '数据库脚本'],
            ['前端开发', 'a', '开发用户聊天界面（Web应用/微信小程序），实现响应式设计，确保界面友好', '前端代码'],
            ['管理员后台', 'a', '开发管理员后台界面', '后台管理代码'],
            ['联调支持', 'b', '协助后端API联调，记录技术实现细节和遇到的挑战与解决方案', '联调记录'],
            ['报告撰写', 'b', '撰写"系统实现"章节', '系统实现文档'],
            ['检查点', '全员', '前后端联调完成', '可运行的完整应用']
        ]
    },
    {
        'title': '第五阶段：集成、测试与部署（第7-9周）',
        'tasks': [
            ['系统集成', 'jiuyong', '前后端及AI服务集成，性能测试和安全测试', '集成测试报告'],
            ['云服务器部署', 'jiuyong', '部署到公有云服务器（阿里云/腾讯云/Heroku等）', '生产环境URL'],
            ['功能测试', 'a', '编写测试用例，进行功能测试、用户体验测试和优化，修复前端bug', '测试用例文档'],
            ['部署配置', 'b', '部署环境配置，准备项目开源代码仓库（GitHub）', 'GitHub仓库'],
            ['报告撰写', 'b', '撰写"系统测试与部署"章节，说明测试用例、部署环境及访问方式', '测试部署文档'],
            ['检查点', '全员', 'Web应用可公开访问，提供URL', '可访问的线上应用']
        ]
    },
    {
        'title': '第六阶段：项目答辩（第9-10周）',
        'tasks': [
            ['报告整合', 'jiuyong', '整合最终报告，审核技术内容，确保技术深度', '最终报告初稿'],
            ['演示准备', 'jiuyong', '准备项目演示脚本，主讲技术实现部分', '演示PPT'],
            ['产品演示', 'a', '准备产品演示流程，展示用户界面和交互功能', '演示流程'],
            ['报告定稿', 'b', '完成最终报告Word文档，确保首页包含Web应用URL和GitHub地址', '最终报告'],
            ['提交报告', 'b', '上传报告到学习通', '学习通提交记录'],
            ['答辩', '全员', '进行最终项目演示与答辩', '答辩完成']
        ]
    }
]

# 添加各阶段
for stage in stages:
    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run(stage['title'])
    set_run_font(run, size=13, bold=True, color=(0, 102, 204))

    # 创建表格
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置表头
    header_cells = table.rows[0].cells
    headers = ['任务', '负责人', '具体工作内容', '交付物']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_font(run, size=10, bold=True, color=(255, 255, 255))
        # 设置背景色
        shading = header_cells[i]._element.get_or_add_tcPr()
        shading_elem = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): '0066CC'
        })
        shading.append(shading_elem)

    # 添加数据行
    for task in stage['tasks']:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(task):
            row_cells[i].text = cell_text
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=10)

# 评分标准
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('评分标准对照')
set_run_font(run, size=14, bold=True)

score_table = doc.add_table(rows=1, cols=4)
score_table.style = 'Table Grid'
score_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# 表头
header_cells = score_table.rows[0].cells
headers = ['评分项', '分值', '主要负责人', '关注重点']
for i, header in enumerate(headers):
    header_cells[i].text = header
    for paragraph in header_cells[i].paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            set_run_font(run, size=10, bold=True, color=(255, 255, 255))
    shading = header_cells[i]._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): '0066CC'
    })
    shading.append(shading_elem)

# 评分数据
scores = [
    ['项目报告', '10%', 'b', '结构完整性、逻辑清晰度、技术深度、文档规范性'],
    ['功能完整性与实用性', '20%', 'jiuyong、a', '核心功能是否完整、是否实用'],
    ['技术实现复杂度与代码质量', '10%', 'jiuyong', '代码质量、技术深度、提示词工程'],
    ['用户界面与交互体验', '10%', 'a', '界面友好、响应式设计、用户体验'],
    ['项目管理与团队协作', '10%', '全员', '更新日志、文档提交、答辩表现']
]

for score in scores:
    row_cells = score_table.add_row().cells
    for i, cell_text in enumerate(score):
        row_cells[i].text = cell_text
        for paragraph in row_cells[i].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size=10)

# 重要提醒
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('重要提醒')
set_run_font(run, size=14, bold=True)

reminders = [
    '提示词工程是评分重点，不仅仅是简单调用API，需要详细记录设计思路和调优过程',
    '每个阶段都要有检查点，确保按时交付',
    '保持代码更新日志，这是项目管理评分的一部分',
    '报告首页必须包含：Web应用URL、GitHub仓库地址'
]

for i, reminder in enumerate(reminders, 1):
    p = doc.add_paragraph(f'{i}. {reminder}')

# 保存
output_path = r'C:\Users\R.westbrook\Desktop\大数据\项目分工表.docx'
doc.save(output_path)
print(f'Word文档已保存到: {output_path}')
