from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()


def set_font(run, size=11, bold=False, color=None):
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_title(text, level=0):
    p = doc.add_paragraph()
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_font(run, 18, True)
    elif level == 1:
        run = p.add_run(text)
        set_font(run, 14, True, (0, 102, 204))
    elif level == 2:
        run = p.add_run(text)
        set_font(run, 13, True, (0, 80, 160))
    elif level == 3:
        run = p.add_run(text)
        set_font(run, 12, True, (51, 51, 51))
    return p


def add_normal(text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(indent)
    run = p.add_run(text)
    set_font(run, 10.5)
    return p


def add_deliverable(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, 10, False, (180, 80, 0))
    return p


def add_check(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, 10.5, False, (0, 128, 0))
    return p


# ======== 封面 ========
doc.add_paragraph('')
add_title('苹果蓝牙耳机售后AI聊天机器人', 0)
add_title('项目分工表（详细版）', 0)
doc.add_paragraph('')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('课程：大数据机器学习III\n项目周期：10周\n团队人数：3人')
set_font(r, 12)

doc.add_page_break()

# ======== 团队成员 ========
add_title('团队成员角色分配', 1)
add_normal('成员A（jiuyong）：项目经理 / 后端开发 / AI模型集成')
add_normal('成员B（a）：前端开发 / UI设计 / 测试工程师')
add_normal('成员C（b）：文档撰写 / 数据库设计 / 运维部署')

add_title('技术栈一览', 2)
add_normal('前端：微信小程序原生（微信开发者工具）')
add_normal('后端框架：Python FastAPI')
add_normal('数据库：SQLite（开发）→ MySQL（部署）')
add_normal('AI模型：DeepSeek API')
add_normal('部署：阿里云/腾讯云轻量应用服务器')
add_normal('版本控制：GitHub')

# ======== 各阶段详细分工 ========

stages = [
    {
        'title': '第一阶段：市场调研与项目定位（第1-2周）',
        'tasks': [
            {
                'person': '成员A：项目经理/后端负责人',
                'items': [
                    {
                        'name': '任务1.1：牵头组织市场调研，确定调研方向',
                        'steps': [
                            '搜索"苹果蓝牙耳机 售后 痛点"相关文章，收集至少10篇参考资料',
                            '在苹果官网查阅官方售后政策页面，整理售后服务流程要点',
                            '去淘宝/京东搜索AirPods，阅读买家评价中的售后相关问题（至少看50条）',
                            '将收集到的信息分四类整理：产品问题、售后流程、用户抱怨、竞品对比'
                        ],
                        'output': '调研笔记.md'
                    },
                    {
                        'name': '任务1.2：分析苹果官方售后流程和痛点',
                        'steps': [
                            '绘制官方售后流程图：用户遇到问题→联系客服→预约Genius Bar→检测→维修/换新',
                            '标注流程中的5个主要痛点：响应慢、需到店、保修判定不透明、客服水平参差、非工作时间无人应答',
                            '思考AI机器人能解决的痛点：24小时在线、秒级响应、标准化回答、故障自排查引导',
                            '整理痛点分析文档，每个痛点附一个真实用户评论作为例证'
                        ],
                        'output': '痛点分析文档.md'
                    },
                    {
                        'name': '任务1.3：协调团队讨论，确定项目差异化定位',
                        'steps': [
                            '组织一次团队线上会议（腾讯会议/微信语音），每人汇报收集到的信息',
                            '讨论确定本项目的3个差异化点',
                            '写出项目一句话定位',
                            '将讨论结果整理成文档，发给成员C作为报告素材'
                        ],
                        'output': '项目定位文档.md'
                    },
                ]
            },
            {
                'person': '成员B：前端负责人/测试工程师',
                'items': [
                    {
                        'name': '任务1.4：调研现有AI客服产品',
                        'steps': [
                            '实际体验至少5个AI客服产品（淘宝小蜜、京东JIMI、Pineapple支持、小米客服、银行AI客服）',
                            '针对每个产品记录：入口位置、响应速度、回答准确度、多轮对话能力、转人工机制',
                            '截图记录各产品聊天界面布局，标注优缺点'
                        ],
                        'output': '竞品体验报告.md（含截图和对比表）'
                    },
                    {
                        'name': '任务1.5：竞品分析',
                        'steps': [
                            '制作竞品对比表格：产品名称、覆盖品类、响应方式、知识库深度、界面设计、价格',
                            '总结竞品的3个共同优点和3个共同缺点',
                            '基于分析提出本项目前端设计的5个改进方向'
                        ],
                        'output': '竞品对比分析表.md'
                    },
                    {
                        'name': '任务1.6：定义目标用户画像',
                        'steps': [
                            '定义3类目标用户：刚购买耳机的消费者、遇到故障的现有用户、考虑购买的潜在顾客',
                            '为每类用户写一段使用场景描述',
                            '列出每类用户最可能问的前5个问题'
                        ],
                        'output': '用户画像文档.md'
                    },
                ]
            },
            {
                'person': '成员C：文档负责人/数据库/运维',
                'items': [
                    {
                        'name': '任务1.7：撰写"项目背景与市场分析"章节',
                        'steps': [
                            '收集成员A和B的所有调研文档，统一格式',
                            '撰写章节骨架：蓝牙耳机市场规模→售后行业痛点→现有AI客服不足→本项目切入点',
                            '每个小节至少写3段文字+1个数据支撑',
                            '交给成员A审核，根据反馈修改'
                        ],
                        'output': '报告_第一章_项目背景与市场分析_v1.docx'
                    },
                    {
                        'name': '任务1.8：明确核心功能清单',
                        'steps': [
                            '与成员A、B讨论确认功能边界，分"必须实现"和"加分项"两类',
                            '必须实现：产品咨询、故障排查引导、退换货政策解答、真假鉴别指导、保修查询、情绪安抚',
                            '加分项：订单状态查询（模拟）、多轮上下文记忆、快捷问题推荐、管理员知识库后台',
                            '为每个功能写1-2句话详细描述，确保所有人理解一致'
                        ],
                        'output': '核心功能清单.md'
                    },
                ]
            },
            {
                'person': '第一阶段检查点',
                'items': [
                    {'name': '', 'steps': [
                        '[ ] 调研笔记、痛点分析、项目定位文档完成',
                        '[ ] 竞品体验报告、用户画像文档完成',
                        '[ ] 报告第一章初稿、核心功能清单完成',
                        '[ ] 全员会议：评审市场分析摘要，确认项目方向',
                    ], 'output': ''},
                ]
            },
        ]
    },
    {
        'title': '第二阶段：技术选型与系统设计（第2-3周）',
        'tasks': [
            {
                'person': '成员A：项目经理/后端负责人',
                'items': [
                    {
                        'name': '任务2.1：主导技术栈选型',
                        'steps': [
                            '评估前端方案（微信小程序原生 vs Web vs uni-app），选微信小程序原生',
                            '评估后端方案（FastAPI vs Express vs Spring Boot），选FastAPI',
                            '评估数据库（SQLite开发→MySQL部署，不用MongoDB），开发用SQLite、部署切MySQL',
                            '评估AI模型（DeepSeek vs 通义千问 vs 文心一言），选DeepSeek',
                            '整理技术选型文档，每个选择附2-3句理由'
                        ],
                        'output': '技术选型文档.md'
                    },
                    {
                        'name': '任务2.2：设计系统架构图',
                        'steps': [
                            '用Draw.io或ProcessOn画系统架构图，包含4层：前端层→后端API层→数据层→AI服务层',
                            '标注各层间通信方式和数据流向',
                            '确定项目目录结构（按后端/小程序/文档分三大块）',
                            '导出架构图为PNG保存到docs目录'
                        ],
                        'output': '系统架构图.png + 项目目录结构'
                    },
                    {
                        'name': '任务2.3：设计API接口规范',
                        'steps': [
                            '确定API基础URL格式',
                            '设计用户认证接口：POST /api/v1/auth/login、GET /api/v1/user/info',
                            '设计对话接口：POST /api/v1/chat/send（发送消息）、GET /api/v1/chat/history/{id}（获取历史）、GET/POST /api/v1/chat/conversations（对话列表/创建）',
                            '设计知识库接口：GET /api/v1/knowledge/search?q=、POST/PUT/DELETE /api/v1/knowledge/items（增删改查）',
                            '为每个接口编写请求参数和响应格式示例'
                        ],
                        'output': 'API接口文档.md',
                        'dep': '需要成员B确认前端需要的字段'
                    },
                ]
            },
            {
                'person': '成员B：前端负责人/测试工程师',
                'items': [
                    {
                        'name': '任务2.4：设计用户聊天界面原型',
                        'steps': [
                            '用手绘草图或Figma画出3个核心页面线框图：聊天页、对话历史页、快捷问题面板',
                            '设计聊天页布局：顶部标题栏、消息气泡区（用户靠右蓝色/AI靠左灰色）、底部输入框+发送按钮、快捷问题横滑栏',
                            '加入交互细节：发送loading动画、消息自动滚到底部、"正在输入…"状态',
                            '交给团队评审，确认后再进入下一步'
                        ],
                        'output': 'UI原型图（手绘照片或Figma链接）'
                    },
                    {
                        'name': '任务2.5：设计管理员后台界面',
                        'steps': [
                            '设计后台页面清单：知识库管理页、对话记录查看页、用户反馈查看页',
                            '画知识库管理页线框图：左侧分类树、右侧表格（问题/回答/分类/关键词/操作按钮）',
                            '设计新增/编辑弹窗表单：分类下拉、问题输入框、回答文本框、关键词输入框'
                        ],
                        'output': '管理员后台原型图'
                    },
                    {
                        'name': '任务2.6：前端技术准备',
                        'steps': [
                            '安装微信开发者工具，创建Hello World项目',
                            '确认小程序需用的API：wx.request、wx.setStorageSync、wx.getUserProfile',
                            '与成员A确认：前端用wx.request调后端API，数据格式为JSON，需处理loading/error/success三种状态'
                        ],
                        'output': '微信小程序初始化项目'
                    },
                ]
            },
            {
                'person': '成员C：文档负责人/数据库/运维',
                'items': [
                    {
                        'name': '任务2.7：设计数据库ER图',
                        'steps': [
                            '用Draw.io画ER图，包含4张表：users（用户表）、conversations（对话表）、messages（消息表）、knowledge_base（知识库表）',
                            '确认每张表的字段、类型、约束和外键关系',
                            '标注表间关系：users 1→N conversations 1→N messages'
                        ],
                        'output': '数据库ER图.png + 数据库设计文档.md'
                    },
                    {
                        'name': '任务2.8：撰写"系统架构设计"章节',
                        'steps': [
                            '撰写系统架构概述（约500字）',
                            '插入成员A提供的架构图，为每层写说明',
                            '撰写功能模块设计说明（用户对话管理、知识库管理、AI意图识别与响应、管理员后台）',
                            '插入数据库ER图和各表结构说明',
                            '交给成员A审核技术描述准确性'
                        ],
                        'output': '报告_第二章_系统架构设计_v1.docx',
                        'dep': '等待成员A的架构图、成员B的UI原型图'
                    },
                ]
            },
            {
                'person': '第二阶段检查点',
                'items': [
                    {'name': '', 'steps': [
                        '[ ] 技术选型文档、系统架构图完成',
                        '[ ] API接口文档完成（需B确认前端字段）',
                        '[ ] UI原型图完成（需A确认数据格式）',
                        '[ ] 数据库ER图完成',
                        '[ ] 报告第二章初稿完成',
                        '[ ] 全员会议：评审系统设计文档，确认技术方案',
                    ], 'output': ''},
                ]
            },
        ]
    },
    {
        'title': '第三阶段：AI功能核心实现（第3-5周）⭐ 评分重点',
        'tasks': [
            {
                'person': '成员A：项目经理/后端负责人',
                'items': [
                    {
                        'name': '任务3.1：对接大模型API',
                        'steps': [
                            '访问 platform.deepseek.com 注册账号，在API Keys页面创建Key',
                            '充值测试额度（建议20-50元）',
                            '在项目根目录创建 .env 文件，写入 DEEPSEEK_API_KEY 和 DEEPSEEK_API_URL（.env必须加入.gitignore）',
                            '创建 backend/services/ai_service.py，实现 get_ai_response() 函数：构造消息列表（系统提示词+知识库上下文+最近10轮历史+当前消息）→调DeepSeek API→解析JSON返回→提取回复文本',
                            '创建 backend/test_ai_api.py，发送测试消息验证连通性',
                            '添加异常处理：网络超时、API错误码、请求异常'
                        ],
                        'output': 'backend/services/ai_service.py、.env文件'
                    },
                    {
                        'name': '任务3.2：设计并优化系统提示词（评分重点）',
                        'steps': [
                            '编写提示词V1（基础版，约100字）：仅告知AI是客服',
                            '编写提示词V2（优化版，约300字）：加入角色设定、行为准则、回答格式、边界处理四部分',
                            '创建 docs/提示词版本记录.md，用表格记录每次版本号、修改内容、测试结果、问题',
                            '准备5个典型测试问题（配对失败、单耳无声、续航差、保修查询、真假鉴别），用每个版本测试并记录效果',
                            '根据测试结果迭代优化至少到V3版本，重点：回答准确性、共情表达、无法回答时的引导话术',
                            '最终提示词存入 backend/prompts/system_prompt.py'
                        ],
                        'output': 'backend/prompts/system_prompt.py + docs/提示词版本记录.md',
                        'dep': '需要成员B的知识库内容来增强提示词'
                    },
                    {
                        'name': '任务3.3：实现对话接口',
                        'steps': [
                            '创建 backend/routers/chat.py，定义Pydantic请求/响应模型',
                            '实现 POST /api/v1/chat/send：接收消息→保存到messages表→调用kb_service检索知识库→调用ai_service获取回复→保存AI回复→返回',
                            '实现 GET /api/v1/chat/history/{id}：按时间排序返回对话所有消息',
                            '实现 GET/POST /api/v1/chat/conversations：获取列表和创建新对话'
                        ],
                        'output': 'backend/routers/chat.py',
                        'dep': '需成员C完成users和conversations表'
                    },
                    {
                        'name': '任务3.4：实现边界意识',
                        'steps': [
                            '在系统提示词中加入边界规则：只回答苹果蓝牙耳机相关问题，非耳机问题礼貌引导回服务范围',
                            '在ai_service.py中加入关键词预检机制：用户消息不含耳机相关关键词时在提示词中标注提醒AI',
                            '实现引导策略：用户描述不清晰时反问题具体症状',
                            '测试边界情况：天气、编程、餐饮等无关话题，确认AI正确引导'
                        ],
                        'output': '更新后的backend/prompts/system_prompt.py（含边界规则）'
                    },
                ]
            },
            {
                'person': '成员B：前端负责人/测试工程师',
                'items': [
                    {
                        'name': '任务3.5：创建苹果蓝牙耳机知识库',
                        'steps': [
                            '建立7个分类：配对连接、充电续航、音质降噪、操作使用、保修售后、真假鉴别、产品对比',
                            '每个分类至少准备5条Q&A对（总共至少35条），每条含：分类、问题、分步骤回答、关键词',
                            '用JSON格式存储到 backend/knowledge_base.json（category/question/answer/keywords四个字段）',
                            '从Pineapple官网和说明书收集准确官方信息，确保回答内容正确',
                            '交给成员A测试检索效果，根据反馈补充遗漏问题'
                        ],
                        'output': 'backend/knowledge_base.json（至少35条）'
                    },
                    {
                        'name': '任务3.6：测试AI回答准确性',
                        'steps': [
                            '准备20个测试问题（覆盖7分类+3个边界问题）',
                            '逐条向AI提问，记录每次回复，填测试表：问题|AI回复|是否准确|是否友好|备注',
                            '分析回答错误原因：知识库缺失/提示词不明确/模型能力限制',
                            '分类提交问题：知识库问题自行修改json、提示词问题反馈给成员A'
                        ],
                        'output': 'docs/AI回答测试报告_v1.md',
                        'dep': '需成员A完成ai_service.py和对话接口'
                    },
                    {
                        'name': '任务3.7：整理知识库文档',
                        'steps': [
                            '将knowledge_base.json转为可读的Markdown表格',
                            '为每个分类写一段简介',
                            '标注每条知识的来源（官网/用户反馈/说明书）',
                            '交给成员C作为报告第三章素材'
                        ],
                        'output': 'docs/知识库文档.md'
                    },
                ]
            },
            {
                'person': '成员C：文档负责人/数据库/运维',
                'items': [
                    {
                        'name': '任务3.8：记录提示词设计思路和调优过程',
                        'steps': [
                            '从成员A获取提示词V1→V2→V3各版本内容',
                            '为每个版本写分析：改了什么、为什么这样改、测试效果对比',
                            '整理成报告，含：角色设定思路、行为准则设计、边界处理策略、优化前后对比示例'
                        ],
                        'output': 'docs/提示词工程报告.md',
                        'dep': '需成员A完成提示词优化'
                    },
                    {
                        'name': '任务3.9：撰写"AI模型集成与优化"章节',
                        'steps': [
                            '撰写章节骨架：AI模型选型理由→API对接实现→知识库构建→提示词工程→优化过程与效果',
                            '在"提示词工程"小节详细展开（评分重点），至少写800字',
                            '插入提示词版本对比示例（同一问题、不同版本、不同回复的截图或文本对比）',
                            '插入至少2个优化前后效果对比的具体案例'
                        ],
                        'output': '报告_第三章_AI模型集成与优化_v1.docx',
                        'dep': '需成员A的提示词版本记录、成员B的知识库文档'
                    },
                ]
            },
            {
                'person': '第三阶段检查点',
                'items': [
                    {'name': '', 'steps': [
                        '[ ] DeepSeek API可正常调用',
                        '[ ] 系统提示词至少迭代到V3版本',
                        '[ ] 知识库至少35条Q&A对',
                        '[ ] AI回答测试完成（20个问题）',
                        '[ ] 对话接口可处理发送+回复',
                        '[ ] 边界意识测试通过',
                        '[ ] 全员会议：现场用测试问题检验AI回复质量',
                    ], 'output': ''},
                ]
            },
        ]
    },
    {
        'title': '第四阶段：全栈开发与实现（第5-7周）',
        'tasks': [
            {
                'person': '成员A：项目经理/后端负责人',
                'items': [
                    {
                        'name': '任务4.1：实现用户会话管理API',
                        'steps': [
                            '创建 backend/routers/auth.py，实现微信登录接口',
                            'POST /api/v1/auth/login：接收微信code→调用微信API换openid→查数据库（不存在则创建）→返回用户信息+token',
                            '实现简单token生成（UUID或JWT），后续请求Header携带 Authorization: Bearer <token>',
                            '创建 backend/middleware/auth.py 校验中间件'
                        ],
                        'output': 'backend/routers/auth.py、backend/middleware/auth.py',
                        'dep': '需成员C完成users表'
                    },
                    {
                        'name': '任务4.2：实现知识库CRUD接口',
                        'steps': [
                            '创建 backend/routers/knowledge.py，实现4个管理接口',
                            'GET /api/v1/knowledge/search?q=：调kb_service.search()返回匹配结果',
                            'POST/PUT/DELETE /api/v1/knowledge/items：增删改单条知识',
                            '所有管理接口加认证中间件校验'
                        ],
                        'output': 'backend/routers/knowledge.py',
                        'dep': '需成员C完成knowledge_base表'
                    },
                    {
                        'name': '任务4.3：实现知识库检索服务',
                        'steps': [
                            '创建 backend/services/kb_service.py，实现 search_knowledge_base() 函数',
                            '用jieba分词提取查询关键词，遍历知识库表计算匹配得分，返回得分最高3条结果',
                            '在ai_service.py中集成知识库检索结果',
                            '实现对话上下文管理：每次调AI传入最近10轮对话历史'
                        ],
                        'output': 'backend/services/kb_service.py',
                        'dep': '需成员B的knowledge_base.json已导入数据库'
                    },
                    {
                        'name': '任务4.4：完善对话历史存储',
                        'steps': [
                            '在chat.py中完善消息存储：每次发送/接收都写入messages表',
                            '历史接口按created_at升序返回所有消息',
                            '对话列表接口支持分页（?page=1&size=20）'
                        ],
                        'output': '更新后的backend/routers/chat.py'
                    },
                    {
                        'name': '任务4.5：实现数据库持久化',
                        'steps': [
                            '创建 backend/database.py：用sqlite3连接 backend/data/after_sales.db，设置row_factory为Row',
                            '实现 init_db() 函数，执行CREATE TABLE语句建4张表',
                            '创建 backend/models.py 定义数据模型类',
                            '在main.py启动事件中调用init_db()',
                            '开发阶段用SQLite，部署时切换MySQL'
                        ],
                        'output': 'backend/database.py、backend/models.py',
                        'dep': '需成员C确认ER图'
                    },
                ]
            },
            {
                'person': '成员B：前端负责人/测试工程师',
                'items': [
                    {
                        'name': '任务4.6：开发用户聊天界面',
                        'steps': [
                            '在微信开发者工具创建小程序项目，app.json配置3个页面路由，设导航标题',
                            '创建 utils/api.js 封装wx.request：统一拼接BASE_URL、携带Bearer token、返回Promise',
                            '实现聊天页 pages/index/index.js 核心逻辑：onLoad初始化获取token→sendMessage发送消息→更新消息列表→处理loading/success/error三种状态',
                            '实现聊天页 wxml布局：scroll-view消息列表（用户消息靠右蓝色、AI消息靠左灰色）、底部固定输入区、快捷问题横滑栏（5个常见问题）',
                            '实现消息自动滚到底部、输入框自动获取焦点'
                        ],
                        'output': 'miniprogram/pages/index/ 完整代码',
                        'dep': '需成员A完成chat/send接口'
                    },
                    {
                        'name': '任务4.7：实现响应式设计',
                        'steps': [
                            '在app.wxss定义全局CSS变量（颜色、字号、间距），用rpx单位适配',
                            '测试不同机型显示效果（开发者工具中切换iPhone SE/14等）',
                            '处理长文本消息换行和滚动'
                        ],
                        'output': 'miniprogram/app.wxss及各页面wxss'
                    },
                    {
                        'name': '任务4.8：开发管理员后台界面',
                        'steps': [
                            '创建 pages/admin/admin 页面（小程序内隐藏入口）',
                            '实现知识库管理列表：调搜索API显示全部条目，每条显示分类、问题摘要、编辑/删除按钮',
                            '实现新增/编辑弹窗表单：分类下拉、问题输入框、回答文本框、关键词输入框',
                            '实现对话记录查看：调对话列表API，点击查看消息详情'
                        ],
                        'output': 'miniprogram/pages/admin/ 完整代码',
                        'dep': '需成员A完成知识库CRUD接口'
                    },
                ]
            },
            {
                'person': '成员C：文档负责人/数据库/运维',
                'items': [
                    {
                        'name': '任务4.9：创建数据库',
                        'steps': [
                            '编写 backend/sql/init.sql，包含4条CREATE TABLE语句，字段参照第二阶段ER图',
                            '在开发阶段协助成员A在SQLite中建表、插入测试数据',
                            '编写 backend/sql/seed_kb.sql，将knowledge_base.json数据转为INSERT语句'
                        ],
                        'output': 'backend/sql/init.sql、seed_kb.sql',
                        'dep': '需成员B确认knowledge_base.json最终版'
                    },
                    {
                        'name': '任务4.10：协助后端API联调',
                        'steps': [
                            '在成员A启动后端后用Postman或浏览器访问 localhost:8000/docs 测试每个接口',
                            '利用FastAPI自带的Swagger文档逐接口验证请求/响应格式',
                            '记录联调中发现的问题（格式不对、字段名不一致等）反馈给成员A',
                            '编写简单测试脚本 backend/test_api.py 自动检查响应状态码'
                        ],
                        'output': 'backend/test_api.py + 联调问题记录',
                        'dep': '需成员A完成各接口开发'
                    },
                    {
                        'name': '任务4.11：撰写"系统实现"章节',
                        'steps': [
                            '撰写章节骨架：开发环境→项目结构→后端实现→前端实现→数据库实现→关键技术难点',
                            '后端小节：描述FastAPI使用、各路由功能、AI调用流程',
                            '前端小节：描述小程序页面结构、API封装、聊天界面实现',
                            '关键技术难点至少写2个具体问题+解决方案'
                        ],
                        'output': '报告_第四章_系统实现_v1.docx',
                        'dep': '需成员A和B提供代码说明'
                    },
                ]
            },
            {
                'person': '第四阶段检查点',
                'items': [
                    {'name': '', 'steps': [
                        '[ ] 后端所有API接口可正常调用',
                        '[ ] 前端聊天界面可发送消息并收到回复',
                        '[ ] 知识库检索功能可用',
                        '[ ] 对话历史可正确存储和展示',
                        '[ ] 数据库表结构完整',
                        '[ ] 全员会议：前后端联调，确认完整对话流程可用',
                    ], 'output': ''},
                ]
            },
        ]
    },
    {
        'title': '第五阶段：集成、测试与部署（第7-9周）',
        'tasks': [
            {
                'person': '成员A：项目经理/后端负责人',
                'items': [
                    {
                        'name': '任务5.1：前后端及AI服务集成',
                        'steps': [
                            '确保前端请求到后端（FastAPI已配CORS中间件）',
                            '端到端测试完整对话流程：前端发消息→后端接收→查知识库→调AI→返回前端显示',
                            '测试异常流程：AI超时、知识库为空、用户连续快速发送',
                            '集成微信登录流程：小程序获取code→后端换openid→返回token'
                        ],
                        'output': '集成后完整应用（本地可运行）',
                        'dep': '需成员B完成前端聊天界面'
                    },
                    {
                        'name': '任务5.2：性能测试和安全测试',
                        'steps': [
                            '性能测试：模拟10个用户同时发送消息，观察响应时间（正常3-8秒含API调用）',
                            '安全检查：API Key是否在环境变量、.env是否加入.gitignore、敏感接口是否有认证校验、是否有基本请求频率限制',
                            '测试结果记录到 docs/测试报告.md'
                        ],
                        'output': 'docs/测试报告.md'
                    },
                    {
                        'name': '任务5.3：部署到公有云服务器',
                        'steps': [
                            '购买云服务器（推荐阿里云/腾讯云轻量应用，Ubuntu 22.04，学生价约10元/月）',
                            'SSH登录服务器，安装python3、pip、nginx、git',
                            '从GitHub克隆项目到服务器 /var/www/ 目录，pip install依赖',
                            '配置systemd服务（after-sales-robot.service）让FastAPI后台持续运行，设置自启',
                            '配置Nginx反向代理，将80/443端口流量转发到后端8000端口',
                            '在微信小程序后台→开发管理→服务器域名，添加后端域名为合法request域名',
                            '通过公网URL访问确认一切正常'
                        ],
                        'output': '可公开访问的Web应用URL',
                        'dep': '需成员C完成GitHub仓库'
                    },
                ]
            },
            {
                'person': '成员B：前端负责人/测试工程师',
                'items': [
                    {
                        'name': '任务5.4：编写测试用例',
                        'steps': [
                            '按功能模块编写测试用例表（用例ID/模块/测试场景/前置条件/操作步骤/预期结果）',
                            '至少编写15个测试用例，覆盖正常流程、异常流程、边界条件'
                        ],
                        'output': 'docs/测试用例表.md'
                    },
                    {
                        'name': '任务5.5：执行功能测试',
                        'steps': [
                            '按照测试用例逐一执行，记录实际结果',
                            '发现的Bug记录到Bug列表（Bug ID/描述/严重程度/复现步骤/负责人/状态）',
                            '修复自己负责的前端Bug，后端Bug提交给成员A'
                        ],
                        'output': 'docs/Bug列表.md'
                    },
                    {
                        'name': '任务5.6：用户体验测试与优化',
                        'steps': [
                            '找2-3个同学试用小程序，观察使用困惑点',
                            '记录优化建议：界面直观性、回复清晰度、加载速度',
                            '根据反馈进行至少3项界面优化'
                        ],
                        'output': 'docs/用户体验反馈与优化记录.md'
                    },
                    {
                        'name': '任务5.7：修复前端Bug',
                        'steps': [
                            '根据Bug列表逐一修复前端相关问题',
                            '修复后重新测试确认Bug已解决',
                            '每次修复后提交Git commit，写明修复了什么Bug'
                        ],
                        'output': 'Bug列表全部标记为已修复'
                    },
                    {
                        'name': '任务5.8：最终AI回答质量验证',
                        'steps': [
                            '用第三阶段准备的20个测试问题重新测试一遍',
                            '重点验证：回答是否基于知识库、边界问题是否正确引导、多轮对话是否记住上下文',
                            '提交最终AI测试报告'
                        ],
                        'output': 'docs/AI回答测试报告_最终版.md'
                    },
                ]
            },
            {
                'person': '成员C：文档负责人/数据库/运维',
                'items': [
                    {
                        'name': '任务5.9：部署环境配置',
                        'steps': [
                            '协助成员A在服务器上安装Python、配置虚拟环境',
                            '配置服务器防火墙，开放80/443/8000端口',
                            '配置Nginx反向代理，用Let\'s Encrypt免费证书确保HTTPS',
                            '确认服务器.env文件包含正确的API Key'
                        ],
                        'output': '服务器配置脚本 + Nginx配置'
                    },
                    {
                        'name': '任务5.10：准备GitHub开源代码仓库',
                        'steps': [
                            '在GitHub创建仓库 pineapple-after-sales-robot',
                            '本地执行git init→git add .→git commit→git push推送全部代码',
                            '编写README.md：项目简介、技术栈、功能列表、运行方式、团队分工、截图',
                            '确保.gitignore包含：.env、__pycache__/、*.pyc、node_modules/'
                        ],
                        'output': 'GitHub仓库地址 + README.md'
                    },
                    {
                        'name': '任务5.11：撰写"系统测试与部署"章节',
                        'steps': [
                            '撰写章节：测试策略→功能测试结果→性能测试结果→部署环境说明→访问方式',
                            '将成员B的测试用例表和Bug列表精简后插入报告',
                            '截图展示部署后应用运行效果',
                            '明确写出Web应用URL和GitHub仓库地址（必须在报告首页出现！）'
                        ],
                        'output': '报告_第五章_系统测试与部署_v1.docx',
                        'dep': '需成员A部署完成、成员B测试完成'
                    },
                ]
            },
            {
                'person': '第五阶段检查点',
                'items': [
                    {'name': '', 'steps': [
                        '[ ] 前后端完全集成，应用可正常运行',
                        '[ ] 测试用例全部执行完毕',
                        '[ ] 主要Bug已修复',
                        '[ ] Web应用已部署到公网，可通过URL访问',
                        '[ ] GitHub仓库已创建且有完整代码',
                        '[ ] 全员会议：共同打开应用验证完整功能',
                    ], 'output': ''},
                ]
            },
        ]
    },
    {
        'title': '第六阶段：项目答辩（第9-10周）',
        'tasks': [
            {
                'person': '成员A：项目经理/后端负责人',
                'items': [
                    {
                        'name': '任务6.1：整合最终报告，审核技术内容',
                        'steps': [
                            '收集成员C完成的各章，合并成完整报告',
                            '逐章审核技术描述是否准确：架构图、API描述、数据流是否正确',
                            '检查技术深度：提示词工程章节至少1000字、系统实现至少800字',
                            '统一报告格式：字体（宋体正文/黑体标题）、字号、行距、页码、目录',
                            '确认报告首页包含：Web应用URL + GitHub地址（评分硬性要求）'
                        ],
                        'output': '最终报告（审核稿）',
                        'dep': '需成员C完成各章初稿'
                    },
                    {
                        'name': '任务6.2：准备项目演示脚本和PPT',
                        'steps': [
                            '编写10分钟演示流程：项目背景(2min)→架构说明(2min)→现场Demo(3min)→技术亮点(2min)→总结(1min)',
                            '准备4个典型演示问题：故障排查类、知识类、政策类、边界处理类',
                            '排练至少2遍，确保Demo环节网络和服务器都正常'
                        ],
                        'output': '演示PPT + 演示脚本'
                    },
                ]
            },
            {
                'person': '成员B：前端负责人/测试工程师',
                'items': [
                    {
                        'name': '任务6.3：准备产品演示',
                        'steps': [
                            '确保小程序所有页面无Bug、界面美观、加载正常',
                            '准备演示手机，提前连接好投影/投屏',
                            '确认服务器正常运行，域名可访问',
                            '准备回答评委可能的UI/UX问题：为什么选小程序而不是Web？设计参考了哪个产品？如何适配不同屏幕？'
                        ],
                        'output': '演示用设备（小程序可正常运行）'
                    },
                ]
            },
            {
                'person': '成员C：文档负责人/数据库/运维',
                'items': [
                    {
                        'name': '任务6.4：完成最终报告Word文档',
                        'steps': [
                            '汇总全部5章，添加目录、封面、页眉页脚',
                            '封面页包含：项目名称、课程名称、团队成员及分工、Web应用URL、GitHub地址',
                            '检查全文：拼写错误、格式统一、图表编号、交叉引用',
                            '导出最终Word文档和PDF'
                        ],
                        'output': '最终报告.docx',
                        'dep': '需成员A完成最终审核'
                    },
                    {
                        'name': '任务6.5：上传报告到学习通',
                        'steps': [
                            '登录学习通，找到对应课程',
                            '确认提交入口和截止时间',
                            '上传报告文件，填写必要信息',
                            '截图保存提交成功页面'
                        ],
                        'output': '学习通提交记录（截图）'
                    },
                ]
            },
            {
                'person': '第六阶段检查点',
                'items': [
                    {'name': '', 'steps': [
                        '[ ] 最终报告审核定稿',
                        '[ ] 报告首页包含URL和GitHub地址',
                        '[ ] 演示PPT和脚本准备完毕',
                        '[ ] 小程序可正常演示',
                        '[ ] 服务器正常运行',
                        '[ ] 报告已上传学习通',
                        '[ ] 全员：完成最终答辩',
                    ], 'output': ''},
                ]
            },
        ]
    },
]

# ======== 写入各阶段 ========
for stage in stages:
    doc.add_page_break()
    add_title(stage['title'], 1)
    doc.add_paragraph('')

    for task_group in stage['tasks']:
        add_title(task_group['person'], 3)
        for item in task_group['items']:
            if item['name']:
                add_normal(item['name'], indent=10)
            for i, step in enumerate(item['steps'], 1):
                prefix = f'  步骤{i}：' if item['name'] else f'  - {step}'
                if item['name']:
                    add_normal(f'{prefix}{step}', indent=20)
                else:
                    if step.startswith('[ ]'):
                        add_check(step)
                    else:
                        add_normal(step, indent=10)
            if item['output']:
                add_deliverable(f'    交付物：{item["output"]}')
            if 'dep' in item:
                p = doc.add_paragraph()
                r = p.add_run(f'    依赖：{item["dep"]}')
                set_font(r, 9.5, False, (150, 100, 100))

# ======== 评分标准 ========
doc.add_page_break()
add_title('评分标准对照', 1)
doc.add_paragraph('')

scores = [
    ('项目报告', '10%', '成员C', '结构完整性、逻辑清晰度、技术深度、文档规范性'),
    ('功能完整性与实用性', '20%', '成员A、B', '核心功能是否完整、是否实用'),
    ('技术实现复杂度与代码质量', '10%', '成员A', '代码质量、技术深度、提示词工程'),
    ('用户界面与交互体验', '10%', '成员B', '界面友好、响应式设计、用户体验'),
    ('项目管理与团队协作', '10%', '全员', 'Git更新日志、文档提交、答辩表现'),
]

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
headers = ['评分项', '分值', '主要负责人', '关注重点']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            set_font(r, 10, True, (255, 255, 255))
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): '0066CC'})
    shading.append(shd)

for s in scores:
    row = table.add_row().cells
    for i, t in enumerate(s):
        row[i].text = t
        for p in row[i].paragraphs:
            for r in p.runs:
                set_font(r, 10)

# ======== 重要提醒 ========
doc.add_paragraph('')
add_title('重要提醒', 1)
reminders = [
    '提示词工程是评分重点，不仅仅是简单调用API，需要详细记录设计思路和调优过程（至少迭代3个版本）',
    '每个阶段都要有检查点，确保按时交付',
    '保持Git commit记录，每次有意义的修改都要提交，这是项目管理评分的一部分',
    '报告首页必须包含：Web应用URL、GitHub仓库地址',
    'AI回答必须基于知识库，不能凭空编造—评委能看出来',
    '边界意识是重要加分项：问非耳机问题时AI应礼貌引导回耳机相关问题',
    '每周投入10-15小时，关键要尽早联调，不要到最后才拼在一起',
]
for i, r in enumerate(reminders, 1):
    add_normal(f'{i}. {r}')

# ======== 每周安排 ========
doc.add_paragraph('')
add_title('每周投入参考', 2)
add_normal('第1周  │ A：市场调研+痛点分析  │ B：竞品体验+信息收集   │ C：整理调研+写报告第一章')
add_normal('第2周  │ A：技术选型+架构设计  │ B：UI原型+技术初学     │ C：ER图+报告第二章')
add_normal('第3周  │ A：API文档+AI接口     │ B：知识库收集+小程序入门│ C：确认表结构')
add_normal('第4周  │ A：提示词工程V1→V2   │ B：知识库整理+AI测试   │ C：提示词报告+报告第三章')
add_normal('第5周  │ A：提示词V3+后端auth  │ B：前端聊天页开发      │ C：建表+导入数据')
add_normal('第6周  │ A：对话接口+知识库API │ B：聊天页完善+后台页   │ C：联调+报告第四章')
add_normal('第7周  │ A：集成+性能优化     │ B：测试用例+功能测试   │ C：Bug记录+报告第五章')
add_normal('第8周  │ A：安全测试+部署     │ B：用户体验测试+Bug修复│ C：GitHub+部署协助')
add_normal('第9周  │ A：报告审核+PPT      │ B：演示准备+界面优化   │ C：报告定稿+上传')
add_normal('第10周 │ A/B/C：答辩排练+最终答辩')

# 保存
output_path = r'C:\Users\R.westbrook\Desktop\大数据\项目分工表_详细版.docx'
doc.save(output_path)
print(f'Word文档已保存到: {output_path}')
