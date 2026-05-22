from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

def set_run_font(run, font_name='微软雅黑', size=11, bold=False, color=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading_styled(text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 0:
        set_run_font(run, size=18, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        set_run_font(run, size=14, bold=True, color=(0, 102, 204))
    elif level == 2:
        set_run_font(run, size=12, bold=True, color=(51, 51, 51))
    return p

def add_step(number, text):
    p = doc.add_paragraph()
    run = p.add_run(f'步骤{number}：')
    set_run_font(run, size=11, bold=True, color=(0, 102, 204))
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p

# 标题
add_heading_styled('苹果蓝牙耳机售后AI聊天机器人项目分工表', 0)

doc.add_paragraph('')

# ========== jiuyong ==========
add_heading_styled('jiuyong（项目经理 / 后端负责人）', 1)

add_heading_styled('一、项目管理', 2)
add_step(1, '牵头组织市场调研，分析苹果官方售后流程和痛点，确定项目差异化定位')
add_step(2, '协调团队讨论，确保各阶段任务按时推进')
add_step(3, '整合最终报告，审核技术内容，确保技术深度')
add_step(4, '准备项目演示脚本，主讲技术实现部分')

add_heading_styled('二、技术选型', 2)
add_step(1, '主导技术栈选型（前端框架、后端框架、数据库、AI模型）')
add_step(2, '设计系统架构，绘制技术架构图')
add_step(3, '设计API接口规范')

add_heading_styled('三、AI核心实现', 2)
add_step(1, '对接大模型API（DeepSeek/通义千问等）')
add_step(2, '实现对话接口，处理用户输入和AI响应')
add_step(3, '设计并优化系统提示词，确保回答专业精准')
add_step(4, '实现边界意识（无法回答时的引导策略）')

add_heading_styled('四、后端开发', 2)
add_step(1, '实现用户会话管理API')
add_step(2, '实现知识库CRUD接口')
add_step(3, '实现AI模型调用代理')
add_step(4, '实现对话历史存储')
add_step(5, '根据ER图创建数据库，实现数据持久化')

add_heading_styled('五、测试与部署', 2)
add_step(1, '前后端及AI服务集成')
add_step(2, '性能测试和安全测试')
add_step(3, '部署到公有云服务器（阿里云/腾讯云/Heroku等）')

doc.add_paragraph('')

# ========== a ==========
add_heading_styled('a（前端负责人 / 测试工程师）', 1)

add_heading_styled('一、市场调研', 2)
add_step(1, '调研现有AI客服机器人的技术方案和用户体验')
add_step(2, '收集竞品信息，分析优劣势')
add_step(3, '协助定义目标用户画像（电商消费者、线下门店顾客等）')

add_heading_styled('二、UI设计', 2)
add_step(1, '设计用户聊天界面原型')
add_step(2, '设计管理员后台界面')

add_heading_styled('三、知识库构建', 2)
add_step(1, '创建苹果蓝牙耳机销售/售后知识库')
add_step(2, '整理Q&A对、产品手册摘要')
add_step(3, '整理退换货政策、常见故障排查内容')

add_heading_styled('四、前端开发', 2)
add_step(1, '开发用户聊天界面（Web应用/微信小程序）')
add_step(2, '实现响应式设计，确保界面友好')
add_step(3, '开发管理员后台界面')

add_heading_styled('五、测试', 2)
add_step(1, '编写测试用例')
add_step(2, '进行功能测试')
add_step(3, '用户体验测试和优化')
add_step(4, '修复前端bug')
add_step(5, '测试AI回答的准确性和专业性，反馈优化建议')

doc.add_paragraph('')

# ========== b ==========
add_heading_styled('b（文档负责人 / 数据库 / 运维）', 1)

add_heading_styled('一、文档撰写', 2)
add_step(1, '整理调研数据，撰写"项目背景与市场分析"章节')
add_step(2, '明确机器人核心功能清单（产品咨询、订单查询、退换货政策解答、故障排查引导、客户情绪安抚）')
add_step(3, '撰写"系统架构设计"章节')
add_step(4, '记录提示词设计思路和调优过程，撰写"AI模型集成与优化"章节')
add_step(5, '整理知识库文档')
add_step(6, '撰写"系统实现"章节，记录技术实现细节和遇到的挑战与解决方案')
add_step(7, '撰写"系统测试与部署"章节')
add_step(8, '完成最终报告Word文档，确保首页包含Web应用URL和GitHub地址')
add_step(9, '上传报告到学习通')

add_heading_styled('二、数据库设计', 2)
add_step(1, '设计数据库ER图')
add_step(2, '规划数据表结构')

add_heading_styled('三、运维部署', 2)
add_step(1, '部署环境配置')
add_step(2, '准备项目开源代码仓库（GitHub）')
add_step(3, '协助后端API联调')

# 保存
output_path = r'C:\Users\R.westbrook\Desktop\大数据\项目分工表_新版.docx'
doc.save(output_path)
print(f'Word文档已保存到: {output_path}')
